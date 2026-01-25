"""
Document metadata extraction module.

Extracts metadata from PDF, Word, Excel, and PowerPoint documents.
Identifies potentially sensitive information like author names, revision history,
and hidden content.
"""

from typing import Optional, Dict, Any, List
from pathlib import Path
from redops.core.context import Context
from redops.core.models import DocumentMetadata, Finding, RiskLevel

# Try to import document libraries
try:
    from pypdf import PdfReader

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Document extensions by type
DOCUMENT_EXTENSIONS = {
    "pdf": {".pdf"},
    "word": {".docx", ".doc"},
    "excel": {".xlsx", ".xls"},
    "powerpoint": {".pptx", ".ppt"},
}

ALL_DOCUMENT_EXTENSIONS = set().union(*DOCUMENT_EXTENSIONS.values())

# Sensitive metadata fields
SENSITIVE_FIELDS = {
    "author": "Document author name",
    "creator": "Creating application/user",
    "last_modified_by": "Last person to modify",
    "manager": "Manager name",
    "company": "Company name",
    "subject": "Document subject",
    "comments": "Document comments",
    "revision": "Revision count (may indicate draft status)",
}


def extract_metadata(ctx: Context, params: Optional[Dict[str, Any]] = None) -> Context:
    """
    Extract metadata from documents.

    Supports PDF, Word (.docx), and scans directories for documents.

    Args:
        ctx: Pipeline context
        params: Optional parameters
            - file_path: Single file to analyze
            - directory: Directory to scan for documents
            - recursive: Scan subdirectories (default: True)
            - include_hidden_check: Check for hidden content (default: True)

    Returns:
        Updated context with document_metadata and findings
    """
    params = params or {}
    file_path = params.get("file_path")
    directory = params.get("directory") or ctx.target
    recursive = params.get("recursive", True)
    include_hidden_check = params.get("include_hidden_check", True)

    if not file_path and not directory:
        ctx.log(
            "No file or directory specified for metadata extraction", level="WARNING"
        )
        return ctx

    metadata_results: List[DocumentMetadata] = []
    findings: List[Finding] = []

    # Process single file
    if file_path:
        ctx.log(f"Extracting metadata from file: {file_path}", level="INFO")
        result = extract_from_file(file_path, include_hidden_check)
        if result:
            metadata_results.append(result)
            file_findings = analyze_document_for_findings(result, file_path)
            findings.extend(file_findings)

    # Process directory
    if directory:
        path = Path(directory)
        if path.is_dir():
            ctx.log(f"Scanning directory for documents: {directory}", level="INFO")
            files = scan_directory_for_documents(directory, recursive=recursive)
            ctx.log(f"Found {len(files)} document files", level="DEBUG")

            for doc_file in files:
                result = extract_from_file(str(doc_file), include_hidden_check)
                if result:
                    metadata_results.append(result)
                    file_findings = analyze_document_for_findings(result, str(doc_file))
                    findings.extend(file_findings)

    # Add to context
    ctx.add("document_metadata", [r.model_dump() for r in metadata_results])

    # Add findings
    for i, finding in enumerate(findings):
        ctx.add(f"finding_doc_{i}", finding.model_dump())

    # Summary statistics
    files_with_author = sum(1 for r in metadata_results if r.author)
    files_with_warnings = sum(1 for r in metadata_results if r.warnings)

    ctx.log(
        f"Metadata extraction completed: {len(metadata_results)} files processed",
        level="INFO",
    )
    if files_with_author > 0:
        ctx.log(
            f"Found {files_with_author} files with author information", level="INFO"
        )
    if files_with_warnings > 0:
        ctx.log(
            f"Found {files_with_warnings} files with hidden content warnings",
            level="WARNING",
        )

    # Add summary
    ctx.add(
        "document_summary",
        {
            "total_files": len(metadata_results),
            "files_with_author": files_with_author,
            "files_with_warnings": files_with_warnings,
            "total_findings": len(findings),
            "by_type": count_by_type(metadata_results),
        },
    )

    return ctx


def scan_directory_for_documents(directory: str, recursive: bool = True) -> List[Path]:
    """
    Scan a directory for document files.

    Args:
        directory: Directory path to scan
        recursive: Whether to scan subdirectories

    Returns:
        List of document file paths
    """
    path = Path(directory)
    document_files = []

    if recursive:
        pattern = "**/*"
    else:
        pattern = "*"

    for file_path in path.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in ALL_DOCUMENT_EXTENSIONS:
            document_files.append(file_path)

    return sorted(document_files)


def extract_from_file(
    file_path: str, include_hidden_check: bool = True
) -> Optional[DocumentMetadata]:
    """
    Extract metadata from a single document file.

    Routes to appropriate extractor based on file type.

    Args:
        file_path: Path to the document
        include_hidden_check: Whether to check for hidden content

    Returns:
        DocumentMetadata object or None
    """
    path = Path(file_path)

    if not path.exists():
        return None

    file_type = path.suffix.lower()

    if file_type not in ALL_DOCUMENT_EXTENSIONS:
        return None

    # Route to appropriate extractor
    if file_type == ".pdf":
        return extract_pdf_metadata(file_path)
    elif file_type in {".docx"}:
        return extract_docx_metadata(file_path, include_hidden_check)
    elif file_type == ".doc":
        # Legacy .doc format not supported without additional libraries
        return DocumentMetadata(
            filename=path.name,
            file_type=file_type,
            metadata={},
            warnings=[
                "Legacy .doc format not fully supported; convert to .docx for analysis"
            ],
        )
    else:
        # Other formats - basic metadata only
        return DocumentMetadata(
            filename=path.name,
            file_type=file_type,
            metadata={"file_size": path.stat().st_size},
            warnings=[f"Full metadata extraction not implemented for {file_type}"],
        )


def extract_pdf_metadata(file_path: str) -> Optional[DocumentMetadata]:
    """
    Extract metadata from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        DocumentMetadata object or None
    """
    path = Path(file_path)

    if not PYPDF_AVAILABLE:
        return DocumentMetadata(
            filename=path.name,
            file_type=".pdf",
            metadata={},
            warnings=["pypdf not available - install with: pip install pypdf"],
        )

    metadata: Dict[str, Any] = {}
    warnings: List[str] = []
    author = None
    created = None
    modified = None

    try:
        reader = PdfReader(file_path)

        # Get document info
        info = reader.metadata
        if info:
            # Standard PDF metadata fields
            if info.author:
                author = str(info.author)
                metadata["author"] = author
            if info.creator:
                metadata["creator"] = str(info.creator)
            if info.producer:
                metadata["producer"] = str(info.producer)
            if info.subject:
                metadata["subject"] = str(info.subject)
            if info.title:
                metadata["title"] = str(info.title)

            # Dates
            if info.creation_date:
                created = str(info.creation_date)
                metadata["created"] = created
            if info.modification_date:
                modified = str(info.modification_date)
                metadata["modified"] = modified

            # Custom metadata
            for key in info:
                if key.startswith("/"):
                    clean_key = key[1:]  # Remove leading slash
                    if clean_key not in metadata:
                        value = info[key]
                        if value:
                            metadata[clean_key] = str(value)

        # Document statistics
        metadata["page_count"] = len(reader.pages)
        metadata["file_size"] = path.stat().st_size

        # Check for encryption
        if reader.is_encrypted:
            warnings.append("Document is encrypted")
            metadata["encrypted"] = True

        # Check for embedded files (potential data exfiltration)
        if reader.attachments:
            attachment_names = list(reader.attachments.keys())
            metadata["attachments"] = attachment_names
            warnings.append(f"Document contains {len(attachment_names)} embedded files")

    except Exception as e:
        warnings.append(f"Error reading PDF: {str(e)}")

    return DocumentMetadata(
        filename=path.name,
        file_type=".pdf",
        author=author,
        created=created,
        modified=modified,
        metadata=metadata,
        warnings=warnings,
    )


def extract_docx_metadata(
    file_path: str, include_hidden_check: bool = True
) -> Optional[DocumentMetadata]:
    """
    Extract metadata from a Word (.docx) file.

    Args:
        file_path: Path to the Word document
        include_hidden_check: Whether to check for hidden content

    Returns:
        DocumentMetadata object or None
    """
    path = Path(file_path)

    if not DOCX_AVAILABLE:
        return DocumentMetadata(
            filename=path.name,
            file_type=".docx",
            metadata={},
            warnings=[
                "python-docx not available - install with: pip install python-docx"
            ],
        )

    metadata: Dict[str, Any] = {}
    warnings: List[str] = []
    author = None
    created = None
    modified = None

    try:
        doc = DocxDocument(file_path)

        # Core properties
        core_props = doc.core_properties

        if core_props.author:
            author = core_props.author
            metadata["author"] = author
        if core_props.last_modified_by:
            metadata["last_modified_by"] = core_props.last_modified_by
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        if core_props.comments:
            metadata["comments"] = core_props.comments
        if core_props.category:
            metadata["category"] = core_props.category
        if core_props.revision:
            metadata["revision"] = core_props.revision

        # Dates
        if core_props.created:
            created = str(core_props.created)
            metadata["created"] = created
        if core_props.modified:
            modified = str(core_props.modified)
            metadata["modified"] = modified

        # Document statistics
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        metadata["file_size"] = path.stat().st_size

        # Check for hidden content if requested
        if include_hidden_check:
            hidden_warnings = check_docx_for_hidden_data(doc)
            warnings.extend(hidden_warnings)

    except PackageNotFoundError:
        warnings.append("Invalid or corrupted Word document")
    except Exception as e:
        warnings.append(f"Error reading Word document: {str(e)}")

    return DocumentMetadata(
        filename=path.name,
        file_type=".docx",
        author=author,
        created=created,
        modified=modified,
        metadata=metadata,
        warnings=warnings,
    )


def check_docx_for_hidden_data(doc) -> List[str]:
    """
    Check a Word document for hidden data.

    Args:
        doc: python-docx Document object

    Returns:
        List of warnings about hidden data
    """
    warnings = []

    try:
        # Check for comments
        # Note: python-docx doesn't have direct comment access in all versions
        # This is a simplified check

        # Check for tracked changes (revisions)
        # This requires checking the XML, which is complex

        # Check for high revision count (indicates many edits)
        if doc.core_properties.revision:
            try:
                rev_count = int(doc.core_properties.revision)
                if rev_count > 10:
                    warnings.append(
                        f"High revision count ({rev_count}) - document has been edited many times"
                    )
            except (ValueError, TypeError):
                pass

        # Check for embedded objects (simplified)
        for rel in doc.part.rels.values():
            if "oleObject" in str(rel.reltype):
                warnings.append("Document contains embedded OLE objects")
                break

    except Exception:
        pass  # Silently ignore errors in hidden data check

    return warnings


def extract_office_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from Office documents (Word, Excel, PowerPoint).

    Generic function that routes to specific extractors.

    Args:
        file_path: Path to the Office document

    Returns:
        Metadata dictionary
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        result = extract_docx_metadata(file_path)
        return result.metadata if result else {}
    else:
        return {
            "filename": path.name,
            "file_type": suffix,
            "warning": f"Extraction for {suffix} not fully implemented",
        }


def check_for_hidden_data(file_path: str) -> List[str]:
    """
    Check for hidden data in documents.

    Args:
        file_path: Path to the document

    Returns:
        List of warnings about hidden data
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx" and DOCX_AVAILABLE:
        try:
            doc = DocxDocument(file_path)
            return check_docx_for_hidden_data(doc)
        except Exception:
            return ["Could not check for hidden data"]

    return []


def analyze_document_for_findings(
    doc_meta: DocumentMetadata, file_path: str
) -> List[Finding]:
    """
    Analyze document metadata and create security findings.

    Args:
        doc_meta: Extracted document metadata
        file_path: Path to the analyzed file

    Returns:
        List of Finding objects
    """
    findings = []

    # Check for author information (medium risk - PII exposure)
    if doc_meta.author:
        finding = Finding(
            module="metadata.documents",
            title=f"Author Information Exposed in {doc_meta.filename}",
            description=f"Document contains author name '{doc_meta.author}'. This may reveal personal or organizational information.",
            severity=RiskLevel.MEDIUM,
            data={
                "file": file_path,
                "author": doc_meta.author,
                "recommendation": "Remove author metadata before sharing externally",
            },
        )
        findings.append(finding)

    # Check for last modified by (low risk)
    last_modified_by = doc_meta.metadata.get("last_modified_by")
    if last_modified_by and last_modified_by != doc_meta.author:
        finding = Finding(
            module="metadata.documents",
            title=f"Editor Information in {doc_meta.filename}",
            description=f"Document was last modified by '{last_modified_by}'.",
            severity=RiskLevel.LOW,
            data={
                "file": file_path,
                "last_modified_by": last_modified_by,
                "recommendation": "Review document properties before distribution",
            },
        )
        findings.append(finding)

    # Check for embedded files/attachments (medium risk)
    if doc_meta.metadata.get("attachments"):
        attachments = doc_meta.metadata["attachments"]
        finding = Finding(
            module="metadata.documents",
            title=f"Embedded Files in {doc_meta.filename}",
            description=f"Document contains {len(attachments)} embedded files which may contain sensitive data.",
            severity=RiskLevel.MEDIUM,
            data={
                "file": file_path,
                "attachments": attachments,
                "recommendation": "Review embedded files for sensitive content",
            },
        )
        findings.append(finding)

    # Check for hidden data warnings (varies by severity)
    for warning in doc_meta.warnings:
        if "embedded" in warning.lower() or "ole" in warning.lower():
            finding = Finding(
                module="metadata.documents",
                title=f"Hidden Content in {doc_meta.filename}",
                description=warning,
                severity=RiskLevel.MEDIUM,
                data={
                    "file": file_path,
                    "warning": warning,
                    "recommendation": "Remove hidden content before sharing",
                },
            )
            findings.append(finding)

    # Check for encryption (info level)
    if doc_meta.metadata.get("encrypted"):
        finding = Finding(
            module="metadata.documents",
            title=f"Encrypted Document: {doc_meta.filename}",
            description="Document is encrypted. Full metadata extraction may be limited.",
            severity=RiskLevel.INFO,
            data={
                "file": file_path,
                "encrypted": True,
            },
        )
        findings.append(finding)

    return findings


def count_by_type(results: List[DocumentMetadata]) -> Dict[str, int]:
    """
    Count documents by file type.

    Args:
        results: List of DocumentMetadata objects

    Returns:
        Dictionary of file type counts
    """
    counts: Dict[str, int] = {}
    for result in results:
        file_type = result.file_type
        counts[file_type] = counts.get(file_type, 0) + 1
    return counts


def get_document_summary(results: List[DocumentMetadata]) -> Dict[str, Any]:
    """
    Generate a summary of document metadata extraction.

    Args:
        results: List of DocumentMetadata objects

    Returns:
        Summary dictionary
    """
    summary = {
        "total_files": len(results),
        "files_with_author": 0,
        "files_with_warnings": 0,
        "by_type": {},
        "authors_found": set(),
        "software_found": set(),
    }

    for result in results:
        if result.author:
            summary["files_with_author"] += 1
            summary["authors_found"].add(result.author)

        if result.warnings:
            summary["files_with_warnings"] += 1

        file_type = result.file_type
        summary["by_type"][file_type] = summary["by_type"].get(file_type, 0) + 1

        if result.metadata.get("creator"):
            summary["software_found"].add(result.metadata["creator"])
        if result.metadata.get("producer"):
            summary["software_found"].add(result.metadata["producer"])

    # Convert sets to lists for JSON serialization
    summary["authors_found"] = list(summary["authors_found"])
    summary["software_found"] = list(summary["software_found"])

    return summary
